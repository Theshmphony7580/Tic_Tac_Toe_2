from docx import Document
import io
import logging

logger = logging.getLogger(__name__)

def parse_docx(file_bytes: bytes) -> str:
    """
    Parses a DOCX file from bytes to plain text, preserving paragraphs and tables.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        
        # Extract Text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract Text from tables
        tables = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    tables.append(" | ".join(row_data))
        
        # Combine
        full_text = "\n".join(paragraphs)
        if tables:
            full_text += "\n\nTables:\n" + "\n".join(tables)
            
        return full_text.strip()
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")
        raise ValueError(f"DOCX parsing failed: {e}")
